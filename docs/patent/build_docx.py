#!/usr/bin/env python3
"""
Build patent DOCX using pandoc only, with proper pipe table syntax.

Strategy:
  1. Use markdown source with pipe tables for first page (proper Word tables)
  2. pandoc converts to DOCX in one shot — proper namespaces, embedded images
"""
import os
import subprocess
from pathlib import Path

PATENT_DIR = Path('/Users/shenggangshu/llm/afk/docs/patent')
MD_FILE = PATENT_DIR / 'AI_Coding_自动化编码工作流助手_技术交底书.md'
OUT_DOCX = PATENT_DIR / 'AI_Coding_自动化编码工作流助手_技术交底书.docx'

# Build the first page as pipe tables (with proper header separators so pandoc recognizes)
FIRST_PAGE = '''| **■ 提案编号**         | AFK-AIC-WF-2026-001        | **■ 提案评级**         | A                  |
|------------------------|------------------------------|------------------------|--------------------|
| **■ 专利工程师**       | 盛纲舒                        | **■ 提案日期**         | 2026 年 08 月 24   |
| **■ 提案名称**         | 基于 AI Coding 的自动化编码工作流助手 ||||
| **■ 所属部门**         | 自动化研发平台部              ||||
| **■ 发明人**           | 盛纲舒、张鹏、石雄飞          |||||
| **■ 联系人（及电话）** | 盛纲舒                        |||||

（请提案人填写）

| **Ⅰ. 本提案的内容所涉及的产品/技术领域是？** |
|---|
| 计算机软件领域中的人工智能辅助开发、DevOps 自动化、多 Agent 协作与工作流编排技术。 |
| **Ⅱ. 本提案的内容所涉及的产品/技术是否已上市、公测或技术公开？** |
| ☒ 否 ☐ 是 （如是，请补充公开方式及日期） |
| **Ⅲ. 就您所知，是否有其他公司使用此技术？** |
| ☒ 否 ☐ 不清楚 ☐ 是 （如是，请列出） |
| **Ⅳ. 本提案的内容在公司内部使用情况：** |
| ☒ 现在有使用且未来仍会使用 ☐ 现在未使用但未来可能会使用 ☐ 现在有使用但未来不会使用 ☐ 现在、未来均不会使用 |
| **Ⅴ. 本提案的内容所涉及的产品是否存在紧急上线的情况？** |
| ☒ 否 ☐ 是 （如是，请列出上线方式及日期） |
| **专利工程师意见：☐ 建议申请 ☐ 建议撤案 ☐ 合并至其他提案 ☐ 待议** |
| **专利评审意见：☐ S ☐ A ☐ B ☐ C ☐ D** |

\\newpage

'''


def main():
    # Strip everything before the body section heading from the source markdown
    body_md = MD_FILE.read_text(encoding='utf-8')
    # Find the position of '## 本发明技术方案的详细阐述' (the actual body start)
    body_start = body_md.find('## 本发明技术方案的详细阐述')
    if body_start < 0:
        raise RuntimeError('Cannot find body start marker')
    body_md = body_md[body_start:]

    # Combine first page + body
    full_md = FIRST_PAGE + '\n\n' + body_md

    # Write to temp file
    tmp_md = PATENT_DIR / '_full.md'
    tmp_md.write_text(full_md, encoding='utf-8')

    # Pandoc with pipe tables
    subprocess.run([
        'pandoc', str(tmp_md), '-o', str(OUT_DOCX),
        '--resource-path=' + str(PATENT_DIR),
        '--from', 'markdown+pipe_tables+raw_html',
    ], check=True)

    os.remove(tmp_md)

    size_kb = OUT_DOCX.stat().st_size / 1024
    print(f'DOCX: {OUT_DOCX} ({size_kb:.1f} KB)')

    # Verify
    import zipfile, re
    with zipfile.ZipFile(OUT_DOCX, 'r') as z:
        files = z.namelist()
        media = [n for n in files if n.startswith('word/media/')]
        doc_xml = z.read('word/document.xml').decode('utf-8')
        rels = z.read('word/_rels/document.xml.rels').decode('utf-8')

    embeds = re.findall(r'r:embed="(rId\d+)"', doc_xml)
    img_rels = re.findall(r'Type="[^"]*image"', rels)
    tables = doc_xml.count('<w:tbl>')
    drawings = doc_xml.count('<w:drawing')

    print(f'\nVerification:')
    print(f'  Tables: {tables} (expected 4: info + qa + metrics + figure-list)')
    print(f'  Drawings: {drawings}')
    print(f'  r:embed: {len(embeds)}')
    print(f'  Image rels: {len(img_rels)}')
    print(f'  Media files: {len(media)}')

    from docx import Document
    try:
        d = Document(str(OUT_DOCX))
        print(f'  python-docx parse: ✓ ({len(d.tables)} tables, {len(d.paragraphs)} paragraphs)')
    except Exception as e:
        print(f'  python-docx parse: ✗ ({e})')


if __name__ == '__main__':
    main()